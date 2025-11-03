import re
import fitz
import os
from pathlib import Path
from llama_index.core import Document, Settings
from llama_index.core.schema import BaseNode
from llama_index.core.node_parser import SentenceSplitter
from dotenv import load_dotenv
from typing import Any, List
from tqdm import tqdm
from ...setting import RAGSettings

load_dotenv()


class LocalDataIngestion:
    def __init__(self, setting: RAGSettings | None = None) -> None:
        self._setting = setting or RAGSettings()
        self._node_store = {}
        self._ingested_file = []

    def _filter_text(self, text):
        # Define the regex pattern.
        pattern = r'[a-zA-Z0-9 \u00C0-\u01B0\u1EA0-\u1EF9`~!@#$%^&*()_\-+=\[\]{}|\\;:\'",.<>/?]+'
        matches = re.findall(pattern, text)
        # Join all matched substrings into a single string
        filtered_text = " ".join(matches)
        # Normalize the text by removing extra whitespaces
        normalized_text = re.sub(r"\s+", " ", filtered_text.strip())

        return normalized_text
    
    def _read_pdf(self, file_path: str) -> list:
        """Read text from PDF file, returns list of (page_num, text) tuples"""
        document = fitz.open(file_path)
        pages_data = []
        
        for page_num, page in enumerate(document, start=1):
            page_text = page.get_text("text")
            page_text = self._filter_text(page_text)
            pages_data.append((page_num, page_text))
        
        document.close()
        
        # Check if we got any text
        total_text = " ".join([text for _, text in pages_data])
        if len(total_text.strip()) < 100:  # Very little text extracted
            print(f"⚠️  Warning: PDF appears to be image-based. Attempting OCR...")
            try:
                # Try OCR if available
                import pytesseract
                from pdf2image import convert_from_path
                from PIL import Image
                
                # Convert PDF pages to images
                images = convert_from_path(file_path)
                pages_data = []
                
                for i, image in enumerate(images, start=1):
                    print(f"   OCR processing page {i}/{len(images)}...")
                    page_text = pytesseract.image_to_string(image)
                    page_text = self._filter_text(page_text)
                    pages_data.append((i, page_text))
                
                print(f"   ✓ OCR completed: {sum(len(t) for _, t in pages_data)} characters extracted")
                
            except ImportError:
                print(f"   ✗ OCR not available. Please install Tesseract OCR.")
                print(f"   See OCR_SETUP.md for installation instructions.")
                print(f"   Returning empty data for this PDF.")
                return []
            except Exception as e:
                print(f"   ✗ OCR failed: {e}")
                print(f"   Returning empty data for this PDF.")
                return []
        
        return pages_data
    
    def _read_txt(self, file_path: str) -> str:
        """Read text from TXT or Markdown file"""
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        return self._filter_text(text)
    
    def _read_docx(self, file_path: str) -> str:
        """Read text from DOCX file"""
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            all_text = ""
            for para in doc.paragraphs:
                all_text += " " + para.text
            return self._filter_text(all_text.strip())
        except ImportError:
            raise ImportError("python-docx is required to read DOCX files. Install with: pip install python-docx")
    
    def _read_file(self, file_path: str) -> str:
        """Read text from file based on extension"""
        ext = Path(file_path).suffix.lower()
        
        if ext == '.pdf':
            return self._read_pdf(file_path)
        elif ext in ['.txt', '.md', '.markdown']:
            return self._read_txt(file_path)
        elif ext == '.docx':
            return self._read_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")

    def store_nodes(
        self,
        input_files: list[str],
        embed_nodes: bool = True,
        embed_model: Any | None = None,
    ) -> List[BaseNode]:
        return_nodes = []
        self._ingested_file = []
        if len(input_files) == 0:
            return return_nodes
        splitter = SentenceSplitter.from_defaults(
            chunk_size=self._setting.ingestion.chunk_size,
            chunk_overlap=self._setting.ingestion.chunk_overlap,
            paragraph_separator=self._setting.ingestion.paragraph_sep,
            secondary_chunking_regex=self._setting.ingestion.chunking_regex,
        )
        if embed_nodes:
            Settings.embed_model = embed_model or Settings.embed_model
        for input_file in tqdm(input_files, desc="Ingesting data"):
            file_name = input_file.strip().split("/")[-1]
            # Also handle Windows paths
            if "\\" in file_name:
                file_name = file_name.split("\\")[-1]
            
            self._ingested_file.append(file_name)
            if file_name in self._node_store:
                return_nodes.extend(self._node_store[file_name])
            else:
                try:
                    # Check file type
                    ext = Path(input_file).suffix.lower()
                    
                    if ext == '.pdf':
                        # Read PDF with page information
                        pages_data = self._read_pdf(input_file)
                        
                        if not pages_data:
                            print(f"No content extracted from {file_name}")
                            continue
                        
                        # Process each page separately to maintain page boundaries
                        nodes = []
                        for page_num, page_text in pages_data:
                            if page_text.strip():  # Only add pages with content
                                doc = Document(
                                    text=page_text,
                                    metadata={
                                        "file_name": file_name,
                                        "page_label": str(page_num),
                                    },
                                )
                                # Split this page into chunks
                                page_nodes = splitter([doc], show_progress=False)
                                # Ensure all chunks from this page have the correct page_label
                                for node in page_nodes:
                                    node.metadata["page_label"] = str(page_num)
                                nodes.extend(page_nodes)
                    else:
                        # Read other file types (returns string)
                        all_text = self._read_file(input_file)
                        
                        document = Document(
                            text=all_text,
                            metadata={
                                "file_name": file_name,
                            },
                        )
                        nodes = splitter([document], show_progress=True)
                    
                    if embed_nodes:
                        nodes = Settings.embed_model(nodes, show_progress=True)
                    self._node_store[file_name] = nodes
                    return_nodes.extend(nodes)
                except Exception as e:
                    print(f"Error processing {file_name}: {e}")
                    continue
        return return_nodes

    def reset(self):
        self._node_store = {}
        self._ingested_file = []

    def check_nodes_exist(self):
        return len(self._node_store.values()) > 0

    def get_all_nodes(self):
        return_nodes = []
        for nodes in self._node_store.values():
            return_nodes.extend(nodes)
        return return_nodes

    def get_ingested_nodes(self):
        return_nodes = []
        for file in self._ingested_file:
            return_nodes.extend(self._node_store[file])
        return return_nodes
